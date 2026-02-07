"""
Script para crear Manual de Administrador en formato Word
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear documento
doc = Document()

# Configurar estilos
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# =======================
# PORTADA
# =======================
title = doc.add_heading('MANUAL DEL ADMINISTRADOR', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Plataforma PQ Trader')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(16)
subtitle_format.font.color.rgb = RGBColor(74, 163, 240)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

version = doc.add_paragraph('Versión 1.0 - Febrero 2026')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
version.runs[0].font.size = Pt(10)
version.runs[0].italic = True

doc.add_page_break()

# =======================
# ÍNDICE
# =======================
doc.add_heading('Tabla de Contenidos', 1)
toc_items = [
    "1. Introducción",
    "2. Primeros Pasos",
    "3. Crear y Gestionar Cursos",
    "4. Agregar Lecciones a los Cursos",
    "5. Subir Archivos y Recursos",
    "6. Publicar Artículos en el Blog",
    "7. Ver Pagos y Transacciones",
    "8. Gestionar Usuarios",
    "9. Solución de Problemas Comunes",
    "10. Preguntas Frecuentes"
]

for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# =======================
# 1. INTRODUCCIÓN
# =======================
doc.add_heading('1. Introducción', 1)

doc.add_paragraph(
    'Este manual le guiará paso a paso para administrar la plataforma PQ Trader. '
    'No necesita conocimientos técnicos previos - todo está explicado de manera simple y clara.'
)

doc.add_heading('¿Qué puede hacer como administrador?', 2)
features = [
    'Crear y publicar cursos de trading',
    'Agregar lecciones con videos y documentos',
    'Escribir y publicar artículos en el blog',
    'Ver quiénes han comprado sus cursos',
    'Revisar los pagos recibidos',
    'Gestionar usuarios de la plataforma'
]

for feature in features:
    p = doc.add_paragraph(f'✓ {feature}', style='List Bullet')

doc.add_heading('Importante antes de comenzar:', 2)
warning = doc.add_paragraph(
    '⚠️ Como administrador, usted tiene acceso completo a la plataforma. '
    'Tenga cuidado al eliminar cursos o usuarios, ya que esta acción no se puede deshacer.'
)
warning.runs[0].font.color.rgb = RGBColor(255, 59, 48)
warning.runs[0].bold = True

doc.add_page_break()

# =======================
# 2. PRIMEROS PASOS
# =======================
doc.add_heading('2. Primeros Pasos', 1)

doc.add_heading('Paso 1: Ingresar a la plataforma', 2)
doc.add_paragraph('1. Abra su navegador web (Chrome, Firefox, Edge, etc.)')
doc.add_paragraph('2. Escriba en la barra de direcciones: https://pqtraders.com/login')
doc.add_paragraph('3. Verá una pantalla de inicio de sesión')

doc.add_heading('Paso 2: Iniciar sesión', 2)
doc.add_paragraph('1. En el campo "Correo electrónico", escriba su email de administrador')
doc.add_paragraph('2. En el campo "Contraseña", escriba su contraseña')
doc.add_paragraph('3. Presione el botón "Iniciar Sesión"')

tip = doc.add_paragraph(
    '💡 Consejo: Si olvidó su contraseña, use el enlace "¿Olvidó su contraseña?" '
    'y recibirá un correo para restablecerla.'
)
tip.runs[0].italic = True

doc.add_heading('Paso 3: Acceder al panel de administración', 2)
doc.add_paragraph('1. Una vez dentro, verá su nombre en la esquina superior derecha')
doc.add_paragraph('2. Haga clic en su nombre')
doc.add_paragraph('3. Aparecerá un menú - seleccione "Mi Dashboard"')
doc.add_paragraph('4. Desde aquí puede acceder a todas las funciones de administración')

doc.add_page_break()

# =======================
# 3. CREAR Y GESTIONAR CURSOS
# =======================
doc.add_heading('3. Crear y Gestionar Cursos', 1)

doc.add_heading('¿Qué es un curso?', 2)
doc.add_paragraph(
    'Un curso es un conjunto de lecciones organizadas que los estudiantes pueden comprar. '
    'Por ejemplo: "Trading Algorítmico con Python" sería un curso que contiene '
    'varias lecciones sobre programación aplicada al trading.'
)

doc.add_heading('Crear un nuevo curso - Paso a Paso', 2)

doc.add_heading('Paso 1: Ir a la sección de cursos', 3)
doc.add_paragraph('1. En el Dashboard, busque el menú lateral izquierdo')
doc.add_paragraph('2. Haga clic en "Cursos"')
doc.add_paragraph('3. Verá la lista de todos los cursos existentes')

doc.add_heading('Paso 2: Crear nuevo curso', 3)
doc.add_paragraph('1. Busque el botón "Crear Nuevo Curso" (usualmente está arriba a la derecha)')
doc.add_paragraph('2. Haga clic en ese botón')
doc.add_paragraph('3. Se abrirá un formulario con varios campos')

doc.add_heading('Paso 3: Llenar la información del curso', 3)

doc.add_paragraph('Ahora debe completar los siguientes campos:')
doc.add_paragraph()

# Tabla de campos
doc.add_paragraph('📝 Título en Español:', style='List Bullet')
doc.add_paragraph('   → Ejemplo: "Trading Algorítmico para Principiantes"')
doc.add_paragraph('   → Este es el nombre que verán los usuarios de habla hispana')

doc.add_paragraph('📝 Título en Inglés:', style='List Bullet')
doc.add_paragraph('   → Ejemplo: "Algorithmic Trading for Beginners"')
doc.add_paragraph('   → Para usuarios que hablen inglés')

doc.add_paragraph('📝 Descripción en Español:', style='List Bullet')
doc.add_paragraph('   → Escriba un texto que explique de qué trata el curso')
doc.add_paragraph('   → Ejemplo: "Aprende a crear tus propias estrategias de trading automatizadas"')

doc.add_paragraph('📝 Descripción en Inglés:', style='List Bullet')
doc.add_paragraph('   → La misma descripción pero en inglés')

doc.add_paragraph('💵 Precio:', style='List Bullet')
doc.add_paragraph('   → Escriba el precio en dólares (solo números)')
doc.add_paragraph('   → Ejemplo: 199.99')
doc.add_paragraph('   → NO escriba el símbolo $, solo el número')

doc.add_paragraph('📊 Nivel:', style='List Bullet')
doc.add_paragraph('   → Seleccione una opción del menú desplegable:')
doc.add_paragraph('      • Principiante - Para personas sin experiencia')
doc.add_paragraph('      • Intermedio - Requiere conocimientos básicos')
doc.add_paragraph('      • Avanzado - Para expertos')

doc.add_paragraph('⏱️ Duración:', style='List Bullet')
doc.add_paragraph('   → Escriba cuántas horas dura el curso completo')
doc.add_paragraph('   → Ejemplo: 25 (solo el número, sin "horas")')

doc.add_paragraph('🖼️ Imagen del curso:', style='List Bullet')
doc.add_paragraph('   → Aquí puede subir una imagen que represente el curso')
doc.add_paragraph('   → Haga clic en "Subir imagen" y seleccione un archivo de su computadora')
doc.add_paragraph('   → Tamaño recomendado: 1200 x 675 píxeles')

doc.add_paragraph('📢 Estado:', style='List Bullet')
doc.add_paragraph('   → Seleccione una opción:')
doc.add_paragraph('      • Borrador - El curso NO será visible para los usuarios')
doc.add_paragraph('      • Publicado - El curso SÍ será visible y se puede comprar')

doc.add_paragraph()
important = doc.add_paragraph(
    '⚠️ IMPORTANTE: Si selecciona "Borrador", el curso se guardará pero los usuarios '
    'NO lo verán hasta que lo cambie a "Publicado".'
)
important.runs[0].font.color.rgb = RGBColor(255, 140, 0)

doc.add_heading('Paso 4: Guardar el curso', 3)
doc.add_paragraph('1. Revise que todos los campos estén correctamente llenados')
doc.add_paragraph('2. Desplácese hasta el final del formulario')
doc.add_paragraph('3. Haga clic en el botón "Crear Curso" o "Guardar"')
doc.add_paragraph('4. Verá un mensaje de confirmación: "Curso creado exitosamente"')

doc.add_heading('Editar un curso existente', 2)
doc.add_paragraph('1. Vaya a "Cursos" en el menú lateral')
doc.add_paragraph('2. Busque el curso que desea modificar')
doc.add_paragraph('3. Haga clic en el icono de lápiz o en el botón "Editar"')
doc.add_paragraph('4. Modifique los campos que necesite')
doc.add_paragraph('5. Haga clic en "Guardar Cambios"')

doc.add_heading('Eliminar un curso', 2)
doc.add_paragraph('1. Vaya a "Cursos"')
doc.add_paragraph('2. Busque el curso que desea eliminar')
doc.add_paragraph('3. Haga clic en el icono de basura o botón "Eliminar"')
doc.add_paragraph('4. Aparecerá un mensaje de confirmación')
doc.add_paragraph('5. Lea el mensaje y confirme si está seguro')

warning2 = doc.add_paragraph(
    '⚠️ CUIDADO: Eliminar un curso es permanente. '
    'Los estudiantes que ya lo compraron no podrán acceder más al contenido. '
    'Si solo quiere ocultarlo temporalmente, mejor cámbielo a "Borrador".'
)
warning2.runs[0].font.color.rgb = RGBColor(255, 59, 48)
warning2.runs[0].bold = True

doc.add_page_break()

# =======================
# 4. AGREGAR LECCIONES
# =======================
doc.add_heading('4. Agregar Lecciones a los Cursos', 1)

doc.add_heading('¿Qué es una lección?', 2)
doc.add_paragraph(
    'Una lección es cada una de las clases o capítulos que forman parte de un curso. '
    'Por ejemplo, si su curso se llama "Trading con Python", podría tener lecciones como:'
)
doc.add_paragraph('   • Lección 1: Introducción a Python', style='List Bullet')
doc.add_paragraph('   • Lección 2: Instalación de librerías', style='List Bullet')
doc.add_paragraph('   • Lección 3: Tu primer script de trading', style='List Bullet')

doc.add_heading('Crear una nueva lección - Paso a Paso', 2)

doc.add_heading('Paso 1: Acceder al curso', 3)
doc.add_paragraph('1. Vaya a "Cursos" en el menú')
doc.add_paragraph('2. Haga clic en el curso donde quiere agregar la lección')
doc.add_paragraph('3. Verá los detalles del curso')

doc.add_heading('Paso 2: Ir a la sección de lecciones', 3)
doc.add_paragraph('1. Dentro del curso, busque la pestaña o sección "Lecciones"')
doc.add_paragraph('2. Haga clic en "Agregar Nueva Lección" o "Crear Lección"')

doc.add_heading('Paso 3: Completar la información de la lección', 3)

doc.add_paragraph('📝 Título de la lección (Español):', style='List Bullet')
doc.add_paragraph('   → Ejemplo: "Introducción a las variables en Python"')

doc.add_paragraph('📝 Título de la lección (Inglés):', style='List Bullet')
doc.add_paragraph('   → Ejemplo: "Introduction to Variables in Python"')

doc.add_paragraph('📄 Contenido (Español):', style='List Bullet')
doc.add_paragraph('   → Aquí escribe el contenido completo de la lección')
doc.add_paragraph('   → Puede usar el editor para dar formato al texto:')
doc.add_paragraph('      • Negritas para resaltar conceptos importantes')
doc.add_paragraph('      • Listas con viñetas para enumerar pasos')
doc.add_paragraph('      • Títulos y subtítulos para organizar el contenido')

doc.add_paragraph('📄 Contenido (Inglés):', style='List Bullet')
doc.add_paragraph('   → El mismo contenido traducido al inglés')

doc.add_paragraph('🎥 URL del Video:', style='List Bullet')
doc.add_paragraph('   → Si tiene un video de la lección en YouTube o Vimeo:')
doc.add_paragraph('      1. Vaya a YouTube y busque su video')
doc.add_paragraph('      2. Copie la URL completa (https://youtube.com/watch?v=...)')
doc.add_paragraph('      3. Péguelo en este campo')
doc.add_paragraph('   → Si no tiene video, puede dejar este campo vacío')

doc.add_paragraph('🔢 Orden:', style='List Bullet')
doc.add_paragraph('   → Escriba un número para indicar la posición de esta lección')
doc.add_paragraph('   → Ejemplo: Si es la primera lección, escriba 1')
doc.add_paragraph('   → Ejemplo: Si es la tercera lección, escriba 3')

doc.add_paragraph('⏱️ Duración:', style='List Bullet')
doc.add_paragraph('   → Escriba cuántos minutos dura esta lección')
doc.add_paragraph('   → Ejemplo: 15 (solo el número)')

doc.add_paragraph('📋 Tipo de lección:', style='List Bullet')
doc.add_paragraph('   → Seleccione qué tipo de lección es:')
doc.add_paragraph('      • Video - Si es principalmente un video')
doc.add_paragraph('      • Texto - Si es un artículo para leer')
doc.add_paragraph('      • Quiz - Si es un cuestionario o examen')
doc.add_paragraph('      • Tarea - Si es un ejercicio para hacer')

doc.add_paragraph('📢 Estado:', style='List Bullet')
doc.add_paragraph('   → Borrador: La lección NO será visible')
doc.add_paragraph('   → Publicado: La lección SÍ será visible')

doc.add_heading('Paso 4: Guardar la lección', 3)
doc.add_paragraph('1. Revise que todo esté correcto')
doc.add_paragraph('2. Haga clic en "Crear Lección" o "Guardar"')
doc.add_paragraph('3. Verá la lección agregada a la lista del curso')

tip2 = doc.add_paragraph(
    '💡 Consejo: Puede crear todas las lecciones en modo "Borrador" primero. '
    'Cuando todo el curso esté listo, cambie todas a "Publicado" al mismo tiempo.'
)
tip2.runs[0].italic = True

doc.add_page_break()

# =======================
# 5. SUBIR ARCHIVOS
# =======================
doc.add_heading('5. Subir Archivos y Recursos', 1)

doc.add_heading('¿Qué archivos puede subir?', 2)
doc.add_paragraph('Puede adjuntar materiales complementarios a sus lecciones:')
doc.add_paragraph('   • Documentos PDF (guías, manuales)', style='List Bullet')
doc.add_paragraph('   • Archivos Excel (.xlsx, .csv)', style='List Bullet')
doc.add_paragraph('   • Código de programación (.py, .r)', style='List Bullet')
doc.add_paragraph('   • Presentaciones PowerPoint', style='List Bullet')
doc.add_paragraph('   • Bases de datos (.csv, .json)', style='List Bullet')

doc.add_heading('Subir un archivo a una lección - Paso a Paso', 2)

doc.add_paragraph('Paso 1: Ir a la lección', style='Heading 3')
doc.add_paragraph('1. Abra el curso que contiene la lección')
doc.add_paragraph('2. Haga clic en la lección donde quiere agregar archivos')
doc.add_paragraph('3. Busque el botón "Editar"')

doc.add_paragraph('Paso 2: Subir archivos', style='Heading 3')
doc.add_paragraph('1. Dentro de la edición de lección, busque la sección "Recursos" o "Archivos adjuntos"')
doc.add_paragraph('2. Haga clic en el botón "Subir archivo" o "Agregar recurso"')
doc.add_paragraph('3. Se abrirá una ventana para seleccionar archivos')
doc.add_paragraph('4. Navegue en su computadora y seleccione el archivo')
doc.add_paragraph('5. Haga clic en "Abrir"')
doc.add_paragraph('6. Espere a que la barra de progreso se complete')
doc.add_paragraph('7. Verá el archivo listado debajo')

doc.add_paragraph('Paso 3: Agregar descripción al archivo', style='Heading 3')
doc.add_paragraph('1. Al lado del archivo subido, puede agregar una descripción')
doc.add_paragraph('2. Ejemplo: "Hoja de ejercicios - Tema 1"')
doc.add_paragraph('3. Esto ayuda a los estudiantes a saber qué es cada archivo')

doc.add_paragraph('Paso 4: Guardar cambios', style='Heading 3')
doc.add_paragraph('1. Haga clic en "Guardar" o "Actualizar lección"')
doc.add_paragraph('2. Los estudiantes podrán descargar los archivos')

important2 = doc.add_paragraph(
    '⚠️ Límite de tamaño: Cada archivo no puede pesar más de 50 MB. '
    'Si tiene archivos más grandes, considere subirlos a Google Drive o Dropbox '
    'y compartir el enlace en el contenido de la lección.'
)
important2.runs[0].font.color.rgb = RGBColor(255, 140, 0)

doc.add_page_break()

# =======================
# 6. BLOG
# =======================
doc.add_heading('6. Publicar Artículos en el Blog', 1)

doc.add_heading('¿Para qué sirve el blog?', 2)
doc.add_paragraph(
    'El blog es una sección donde puede publicar artículos educativos gratuitos. '
    'Estos artículos ayudan a:'
)
doc.add_paragraph('   • Atraer nuevos estudiantes con contenido de valor', style='List Bullet')
doc.add_paragraph('   • Posicionar su marca como experto en trading', style='List Bullet')
doc.add_paragraph('   • Mejorar el SEO (aparecer en Google)', style='List Bullet')
doc.add_paragraph('   • Compartir noticias y novedades', style='List Bullet')

doc.add_heading('Crear un artículo - Paso a Paso', 2)

doc.add_paragraph('Paso 1: Ir a Blog', style='Heading 3')
doc.add_paragraph('1. En el menú lateral, haga clic en "Blog"')
doc.add_paragraph('2. Verá todos los artículos publicados')
doc.add_paragraph('3. Haga clic en "Crear Nuevo Artículo"')

doc.add_paragraph('Paso 2: Completar la información', style='Heading 3')

doc.add_paragraph('📝 Título del artículo (Español):', style='List Bullet')
doc.add_paragraph('   → Ejemplo: "10 Consejos para el Trading Algorítmico"')
doc.add_paragraph('   → Use un título atractivo que capte atención')

doc.add_paragraph('📝 Título del artículo (Inglés):', style='List Bullet')
doc.add_paragraph('   → Traducción del título')

doc.add_paragraph('🔗 Slug (URL amigable):', style='List Bullet')
doc.add_paragraph('   → Este campo se genera automáticamente')
doc.add_paragraph('   → Ejemplo: "10-consejos-trading-algoritmico"')
doc.add_paragraph('   → Será la dirección web del artículo')

doc.add_paragraph('✍️ Contenido del artículo:', style='List Bullet')
doc.add_paragraph('   → Use el editor para escribir el artículo completo')
doc.add_paragraph('   → Puede agregar:')
doc.add_paragraph('      • Títulos y subtítulos para organizar')
doc.add_paragraph('      • Negritas y cursivas para énfasis')
doc.add_paragraph('      • Listas numeradas o con viñetas')
doc.add_paragraph('      • Enlaces a otros sitios web')
doc.add_paragraph('      • Imágenes (haga clic en el icono de imagen)')

doc.add_paragraph('📄 Extracto o resumen:', style='List Bullet')
doc.add_paragraph('   → Escriba un resumen corto del artículo (2-3 líneas)')
doc.add_paragraph('   → Este texto aparecerá en la lista de artículos')

doc.add_paragraph('🖼️ Imagen destacada:', style='List Bullet')
doc.add_paragraph('   → Suba una imagen representativa del artículo')
doc.add_paragraph('   → Aparecerá en la portada del artículo')
doc.add_paragraph('   → Tamaño recomendado: 1200 x 630 píxeles')

doc.add_paragraph('🏷️ Categorías:', style='List Bullet')
doc.add_paragraph('   → Seleccione las categorías que apliquen:')
doc.add_paragraph('      • Trading - Para artículos sobre estrategias')
doc.add_paragraph('      • Python - Para programación')
doc.add_paragraph('      • Análisis - Para análisis técnico o fundamental')
doc.add_paragraph('      • Principiantes - Para contenido introductorio')

doc.add_paragraph('⭐ Artículo destacado:', style='List Bullet')
doc.add_paragraph('   → Active esta casilla si quiere que aparezca en la página principal')
doc.add_paragraph('   → Solo marque los mejores artículos como destacados')

doc.add_paragraph('📢 Estado:', style='List Bullet')
doc.add_paragraph('   → Borrador - El artículo NO será visible')
doc.add_paragraph('   → Publicado - El artículo SÍ será visible en el blog')

doc.add_paragraph('Paso 3: Previsualizar', style='Heading 3')
doc.add_paragraph('1. Antes de publicar, puede ver cómo se verá el artículo')
doc.add_paragraph('2. Busque el botón "Previsualizar"')
doc.add_paragraph('3. Se abrirá una nueva pestaña mostrando el artículo')
doc.add_paragraph('4. Si necesita hacer cambios, regrese y edite')

doc.add_paragraph('Paso 4: Publicar', style='Heading 3')
doc.add_paragraph('1. Cuando esté satisfecho con el artículo:')
doc.add_paragraph('2. Cambie el estado a "Publicado"')
doc.add_paragraph('3. Haga clic en "Publicar Artículo"')
doc.add_paragraph('4. El artículo estará disponible instantáneamente')

tip3 = doc.add_paragraph(
    '💡 Consejo: Publique artículos regularmente (1-2 por semana) para mantener '
    'a su audiencia interesada y mejorar su posicionamiento en buscadores.'
)
tip3.runs[0].italic = True

doc.add_page_break()

# =======================
# 7. PAGOS Y TRANSACCIONES
# =======================
doc.add_heading('7. Ver Pagos y Transacciones', 1)

doc.add_heading('¿Dónde veo los pagos recibidos?', 2)
doc.add_paragraph('1. En el menú lateral, haga clic en "Finanzas" o "Transacciones"')
doc.add_paragraph('2. Verá una tabla con todos los pagos')

doc.add_heading('Información que verá en cada transacción:', 2)
doc.add_paragraph('   • Fecha y hora del pago', style='List Bullet')
doc.add_paragraph('   • Nombre del estudiante', style='List Bullet')
doc.add_paragraph('   • Producto comprado (curso o mentoría)', style='List Bullet')
doc.add_paragraph('   • Monto pagado en dólares', style='List Bullet')
doc.add_paragraph('   • Método de pago (Stripe o PayPal)', style='List Bullet')
doc.add_paragraph('   • Estado del pago', style='List Bullet')

doc.add_heading('Estados de pago:', 2)
states_table = doc.add_table(rows=5, cols=2)
states_table.style = 'Light Grid Accent 1'

# Headers
hdr_cells = states_table.rows[0].cells
hdr_cells[0].text = 'Estado'
hdr_cells[1].text = 'Significado'

# Rows
row_data = [
    ('Pendiente', 'El pago está en proceso'),
    ('Completado', 'El pago fue exitoso'),
    ('Fallido', 'El pago fue rechazado'),
    ('Reembolsado', 'Se devolvió el dinero al cliente')
]

for i, (estado, significado) in enumerate(row_data, start=1):
    row_cells = states_table.rows[i].cells
    row_cells[0].text = estado
    row_cells[1].text = significado

doc.add_paragraph()

doc.add_heading('Filtrar transacciones:', 2)
doc.add_paragraph('1. Use los filtros en la parte superior de la tabla')
doc.add_paragraph('2. Puede filtrar por:')
doc.add_paragraph('   • Fecha (última semana, último mes, rango personalizado)', style='List Bullet')
doc.add_paragraph('   • Estado (completado, pendiente, etc.)', style='List Bullet')
doc.add_paragraph('   • Método de pago (Stripe, PayPal)', style='List Bullet')

doc.add_heading('Descargar reporte:', 2)
doc.add_paragraph('1. Seleccione el período que desea')
doc.add_paragraph('2. Haga clic en "Exportar a Excel" o "Descargar CSV"')
doc.add_paragraph('3. Se descargará un archivo con todas las transacciones')
doc.add_paragraph('4. Puede abrirlo en Excel para hacer análisis')

doc.add_heading('Procesar reembolsos:', 2)
doc.add_paragraph(
    'Si necesita devolver el dinero a un cliente, NO lo haga desde la plataforma PQ Trader. '
    'Debe hacerlo directamente desde Stripe o PayPal:'
)

doc.add_paragraph('Para reembolsos en Stripe:', style='Heading 3')
doc.add_paragraph('1. Vaya a https://dashboard.stripe.com')
doc.add_paragraph('2. Inicie sesión con su cuenta de Stripe')
doc.add_paragraph('3. Haga clic en "Pagos" en el menú izquierdo')
doc.add_paragraph('4. Busque la transacción que desea reembolsar')
doc.add_paragraph('5. Haga clic en ella')
doc.add_paragraph('6. Arriba a la derecha verá el botón "Refund"')
doc.add_paragraph('7. Ingrese el monto a reembolsar (puede ser parcial)')
doc.add_paragraph('8. Haga clic en "Refund"')
doc.add_paragraph('9. El dinero se devolverá en 5-10 días hábiles')

doc.add_paragraph('Para reembolsos en PayPal:', style='Heading 3')
doc.add_paragraph('1. Vaya a https://www.paypal.com')
doc.add_paragraph('2. Inicie sesión')
doc.add_paragraph('3. Haga clic en "Actividad"')
doc.add_paragraph('4. Busque la transacción')
doc.add_paragraph('5. Haga clic en ella')
doc.add_paragraph('6. Haga clic en "Reembolsar este pago"')
doc.add_paragraph('7. Ingrese el monto')
doc.add_paragraph('8. Haga clic en "Enviar reembolso"')

important3 = doc.add_paragraph(
    '⚠️ IMPORTANTE: Después de hacer el reembolso en Stripe o PayPal, '
    'el estado en PQ Trader se actualizará automáticamente a "Reembolsado".'
)
important3.runs[0].font.color.rgb = RGBColor(255, 140, 0)

doc.add_page_break()

# =======================
# 8. GESTIÓN DE USUARIOS
# =======================
doc.add_heading('8. Gestionar Usuarios', 1)

doc.add_heading('Ver todos los usuarios:', 2)
doc.add_paragraph('1. En el menú lateral, haga clic en "Usuarios"')
doc.add_paragraph('2. Verá una lista con todos los usuarios registrados')
doc.add_paragraph('3. Para cada usuario verá:')
doc.add_paragraph('   • Nombre completo', style='List Bullet')
doc.add_paragraph('   • Correo electrónico', style='List Bullet')
doc.add_paragraph('   • Fecha de registro', style='List Bullet')
doc.add_paragraph('   • Rol (Estudiante o Administrador)', style='List Bullet')
doc.add_paragraph('   • Cursos comprados', style='List Bullet')

doc.add_heading('Buscar un usuario específico:', 2)
doc.add_paragraph('1. En la parte superior de la lista, verá una barra de búsqueda')
doc.add_paragraph('2. Escriba el nombre o correo del usuario')
doc.add_paragraph('3. La lista se filtrará automáticamente')

doc.add_heading('Ver detalles de un usuario:', 2)
doc.add_paragraph('1. Haga clic en el nombre del usuario')
doc.add_paragraph('2. Se abrirá una página con información detallada:')
doc.add_paragraph('   • Datos personales', style='List Bullet')
doc.add_paragraph('   • Historial de compras', style='List Bullet')
doc.add_paragraph('   • Progreso en los cursos', style='List Bullet')
doc.add_paragraph('   • Total gastado', style='List Bullet')

doc.add_heading('Cambiar el rol de un usuario:', 2)
doc.add_paragraph(
    'Puede dar permisos de administrador a otro usuario (por ejemplo, un asistente):'
)
doc.add_paragraph('1. Vaya a la página del usuario')
doc.add_paragraph('2. Busque la sección "Rol"')
doc.add_paragraph('3. Haga clic en "Cambiar rol"')
doc.add_paragraph('4. Seleccione "Administrador" o "Estudiante"')
doc.add_paragraph('5. Confirme el cambio')

warning3 = doc.add_paragraph(
    '⚠️ CUIDADO: Un administrador tiene acceso COMPLETO a la plataforma. '
    'Solo otorgue este rol a personas de confianza.'
)
warning3.runs[0].font.color.rgb = RGBColor(255, 59, 48)
warning3.runs[0].bold = True

doc.add_heading('Eliminar un usuario:', 2)
doc.add_paragraph('1. Vaya a la página del usuario')
doc.add_paragraph('2. Desplácese hasta el final')
doc.add_paragraph('3. Haga clic en "Eliminar usuario"')
doc.add_paragraph('4. Aparecerá una advertencia')
doc.add_paragraph('5. Confirme si está seguro')

important4 = doc.add_paragraph(
    '⚠️ Al eliminar un usuario, se borrarán todos sus datos y NO se puede deshacer. '
    'El usuario perderá acceso a sus cursos comprados.'
)
important4.runs[0].font.color.rgb = RGBColor(255, 140, 0)

doc.add_page_break()

# =======================
# 9. SOLUCIÓN DE PROBLEMAS
# =======================
doc.add_heading('9. Solución de Problemas Comunes', 1)

doc.add_heading('Problema 1: No puedo iniciar sesión', 2)
doc.add_paragraph('Posibles causas:', style='Heading 3')
doc.add_paragraph('   • Contraseña incorrecta', style='List Bullet')
doc.add_paragraph('   • Email mal escrito', style='List Bullet')
doc.add_paragraph('   • Cuenta bloqueada', style='List Bullet')

doc.add_paragraph('Solución:', style='Heading 3')
doc.add_paragraph('1. Verifique que está escribiendo el email correctamente')
doc.add_paragraph('2. Pruebe con "Recuperar contraseña"')
doc.add_paragraph('3. Revise su correo (también la carpeta Spam)')
doc.add_paragraph('4. Si persiste, contacte al soporte técnico')

doc.add_paragraph()
doc.add_heading('Problema 2: No veo el curso que acabo de crear', 2)
doc.add_paragraph('Posibles causas:', style='Heading 3')
doc.add_paragraph('   • El curso está en modo "Borrador"', style='List Bullet')
doc.add_paragraph('   • No se guardó correctamente', style='List Bullet')

doc.add_paragraph('Solución:', style='Heading 3')
doc.add_paragraph('1. Vaya a la lista de cursos')
doc.add_paragraph('2. Busque su curso')
doc.add_paragraph('3. Verifique que el estado sea "Publicado"')
doc.add_paragraph('4. Si está en "Borrador", edítelo y cámbielo a "Publicado"')

doc.add_paragraph()
doc.add_heading('Problema 3: Un archivo no se sube', 2)
doc.add_paragraph('Posibles causas:', style='Heading 3')
doc.add_paragraph('   • El archivo es muy pesado (más de 50 MB)', style='List Bullet')
doc.add_paragraph('   • Conexión a internet lenta', style='List Bullet')
doc.add_paragraph('   • Formato de archivo no permitido', style='List Bullet')

doc.add_paragraph('Solución:', style='Heading 3')
doc.add_paragraph('1. Verifique el tamaño del archivo (clic derecho → Propiedades)')
doc.add_paragraph('2. Si es muy grande, comprímalo o divídalo en partes')
doc.add_paragraph('3. Intente con una conexión a internet más rápida')
doc.add_paragraph('4. Verifique que sea un formato permitido (PDF, Excel, etc.)')

doc.add_paragraph()
doc.add_heading('Problema 4: Un usuario dice que no recibió acceso después de pagar', 2)
doc.add_paragraph('Solución:', style='Heading 3')
doc.add_paragraph('1. Vaya a "Finanzas" → "Transacciones"')
doc.add_paragraph('2. Busque la transacción por el nombre del usuario o email')
doc.add_paragraph('3. Verifique que el estado sea "Completado"')
doc.add_paragraph('4. Si está en "Pendiente", espere unos minutos')
doc.add_paragraph('5. Si está en "Fallido", el pago no se procesó')
doc.add_paragraph('6. Pídale al usuario que intente nuevamente')
doc.add_paragraph('7. Si el problema persiste, contacte a soporte técnico')

doc.add_paragraph()
doc.add_heading('Problema 5: Olvidé mi contraseña', 2)
doc.add_paragraph('Solución:', style='Heading 3')
doc.add_paragraph('1. Vaya a https://pqtraders.com/login')
doc.add_paragraph('2. Haga clic en "¿Olvidó su contraseña?"')
doc.add_paragraph('3. Ingrese su correo electrónico')
doc.add_paragraph('4. Haga clic en "Enviar"')
doc.add_paragraph('5. Recibirá un email con instrucciones')
doc.add_paragraph('6. Revise también su carpeta de Spam')
doc.add_paragraph('7. Siga el enlace del email para crear una nueva contraseña')

doc.add_paragraph()
doc.add_heading('Problema 6: El video de una lección no se reproduce', 2)
doc.add_paragraph('Solución:', style='Heading 3')
doc.add_paragraph('1. Verifique que la URL del video sea correcta')
doc.add_paragraph('2. Abra el enlace en una nueva pestaña para probarlo')
doc.add_paragraph('3. Si el video está en YouTube, asegúrese que NO sea privado')
doc.add_paragraph('4. El video debe estar configurado como "Público" o "No listado"')
doc.add_paragraph('5. Si es privado, los estudiantes no podrán verlo')

doc.add_page_break()

# =======================
# 10. PREGUNTAS FRECUENTES
# =======================
doc.add_heading('10. Preguntas Frecuentes', 1)

doc.add_heading('¿Puedo editar un curso después de publicarlo?', 2)
doc.add_paragraph(
    'Sí, puede editar cursos en cualquier momento. Los cambios se reflejarán inmediatamente '
    'para todos los estudiantes que ya lo compraron.'
)

doc.add_heading('¿Cuándo recibo el dinero de las ventas?', 2)
doc.add_paragraph('Pagos con Stripe:', style='Heading 3')
doc.add_paragraph(
    'El dinero se transfiere automáticamente a su cuenta bancaria cada 2-7 días hábiles, '
    'dependiendo de la configuración de su cuenta Stripe.'
)

doc.add_paragraph('Pagos con PayPal:', style='Heading 3')
doc.add_paragraph(
    'El dinero está disponible inmediatamente en su cuenta PayPal. '
    'Puede transferirlo a su banco cuando lo desee.'
)

doc.add_heading('¿Puedo ofrecer cupones de descuento?', 2)
doc.add_paragraph(
    'Actualmente esta funcionalidad está en desarrollo. '
    'Por ahora, si desea ofrecer un descuento, puede editar temporalmente el precio del curso.'
)

doc.add_heading('¿Los estudiantes pueden descargar los videos?', 2)
doc.add_paragraph(
    'Depende de la configuración de su cuenta de YouTube o Vimeo. '
    'Por seguridad, se recomienda NO permitir descargas para proteger su contenido.'
)

doc.add_heading('¿Puedo tener varios administradores?', 2)
doc.add_paragraph(
    'Sí, puede otorgar permisos de administrador a otros usuarios. '
    'Vaya a "Usuarios" → seleccione el usuario → "Cambiar rol" → "Administrador".'
)

doc.add_heading('¿Cómo sé si un estudiante completó un curso?', 2)
doc.add_paragraph(
    'Vaya a "Usuarios" → seleccione el estudiante → verá su progreso en cada curso, '
    'indicando qué lecciones ha completado y su porcentaje de avance.'
)

doc.add_heading('¿Puedo borrar lecciones de un curso publicado?', 2)
doc.add_paragraph(
    'Sí, pero tenga cuidado: si elimina una lección, los estudiantes que ya compraron '
    'el curso perderán acceso a ese contenido. Es mejor editar la lección que eliminarla.'
)

doc.add_heading('¿Cómo cambio el precio de un curso?', 2)
doc.add_paragraph('1. Vaya a "Cursos"')
doc.add_paragraph('2. Haga clic en "Editar" en el curso')
doc.add_paragraph('3. Modifique el campo "Precio"')
doc.add_paragraph('4. Haga clic en "Guardar Cambios"')
doc.add_paragraph(
    'Nota: El nuevo precio solo aplicará para nuevas compras. '
    'Los estudiantes que ya compraron mantendrán su acceso.'
)

doc.add_heading('¿Qué hago si un estudiante tiene problemas técnicos?', 2)
doc.add_paragraph('1. Pídale que le envíe capturas de pantalla del error')
doc.add_paragraph('2. Verifique que el curso y las lecciones estén publicados')
doc.add_paragraph('3. Revise que el pago se haya completado en "Transacciones"')
doc.add_paragraph('4. Si todo está bien, contacte a soporte técnico: info@pqtraders.com')

doc.add_page_break()

# =======================
# INFORMACIÓN DE CONTACTO
# =======================
doc.add_heading('Información de Contacto', 1)

doc.add_heading('Soporte Técnico', 2)
contact_info = [
    ('Email:', 'info@pqtraders.com'),
    ('Sitio web:', 'https://pqtraders.com'),
    ('Horario de atención:', 'Lunes a Viernes, 9:00 AM - 6:00 PM (GMT-5)'),
    ('Tiempo de respuesta:', '24-48 horas hábiles')
]

for label, value in contact_info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(f' {value}')

doc.add_paragraph()
important5 = doc.add_paragraph(
    '💡 Consejo: Al contactar a soporte, incluya siempre capturas de pantalla '
    'y una descripción detallada del problema. Esto ayudará a resolver su consulta más rápidamente.'
)
important5.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

doc.add_heading('Accesos Importantes', 2)
access_info = [
    ('Plataforma:', 'https://pqtraders.com'),
    ('Panel de administración:', 'https://pqtraders.com/dashboard'),
    ('Stripe Dashboard:', 'https://dashboard.stripe.com'),
    ('PayPal:', 'https://www.paypal.com')
]

for label, value in access_info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(f' {value}')

doc.add_page_break()

# =======================
# GLOSARIO
# =======================
doc.add_heading('Glosario de Términos', 1)

doc.add_paragraph(
    'Aquí encontrará el significado de términos técnicos que aparecen en este manual:'
)

doc.add_paragraph()

glossary = [
    ('Blog', 'Sección de artículos educativos en el sitio web'),
    ('Borrador', 'Estado de un curso/lección/artículo que NO es visible para los usuarios'),
    ('Dashboard', 'Panel de control donde administra todo el contenido'),
    ('Lección', 'Cada clase o capítulo que forma parte de un curso'),
    ('PayPal', 'Plataforma de pagos en línea que permite recibir dinero de clientes'),
    ('Publicado', 'Estado que hace visible un curso/lección/artículo a los usuarios'),
    ('Reembolso', 'Devolución de dinero a un cliente'),
    ('Slug', 'Versión simplificada del título usada en la URL (ejemplo: mi-primer-curso)'),
    ('Stripe', 'Plataforma de pagos que procesa tarjetas de crédito y débito'),
    ('Transacción', 'Cada pago realizado en la plataforma'),
    ('URL', 'Dirección web (ejemplo: https://pqtraders.com/cursos)'),
    ('Video URL', 'Enlace directo a un video en YouTube, Vimeo u otra plataforma')
]

for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(f'{term}: ').bold = True
    p.add_run(definition)

doc.add_paragraph()
doc.add_paragraph()

# =======================
# FIN DEL DOCUMENTO
# =======================
final_note = doc.add_paragraph()
final_note.add_run('Este manual fue creado en ').italic = True
final_note.add_run('Febrero 2026').italic = True
final_note.add_run(' para la plataforma PQ Trader.').italic = True
final_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

final_note2 = doc.add_paragraph('Si tiene sugerencias para mejorar este manual, escríbanos a info@pqtraders.com')
final_note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_note2.runs[0].font.size = Pt(9)

# Guardar documento
doc.save('c:/Users/USER/Desktop/pq_trader/docs/Manual_Administrador_PQ_Trader.docx')
print("✅ Manual creado exitosamente: Manual_Administrador_PQ_Trader.docx")
